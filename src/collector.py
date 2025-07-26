import os
import re
import json
import time
import logging
import requests
from pathlib import Path
from typing import Dict, List, Any, Optional, Union
from urllib.parse import urljoin, urlparse
from datetime import datetime
import hashlib
import mimetypes
import io
import os 
import sys

# fmt: off
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
sys.path.append(os.path.dirname(os.path.dirname(SCRIPT_DIR)))

# Try to import custom logger, fallback to standard logging
try:
    from logger.custom_logger import CustomLoggerTracker
    logger_tracker = CustomLoggerTracker()
    logger = logger_tracker.get_logger("main")
    logger.info("Custom logger initialized")
except ImportError:
    import logging
    logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
    logger = logging.getLogger("main")
    logger.info("Using standard logger - custom logger not available")



# Web scraping dependencies
try:
    from bs4 import BeautifulSoup
    import requests
    WEB_SCRAPING_AVAILABLE = True
except ImportError:
    print("Warning: Web scraping dependencies not available. Install with: pip install beautifulsoup4 requests")
    WEB_SCRAPING_AVAILABLE = False

# PDF processing dependencies
try:
    import PyPDF2
    import pdfplumber
    PDF_PROCESSING_AVAILABLE = True
except ImportError:
    print("Warning: PDF processing dependencies not available. Install with: pip install PyPDF2 pdfplumber")
    PDF_PROCESSING_AVAILABLE = False

# Document processing dependencies
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Additional text processing
try:
    import nltk
    from nltk.tokenize import sent_tokenize, word_tokenize
    NLTK_AVAILABLE = True
except ImportError:
    NLTK_AVAILABLE = False




class DataCollector:
    """
    Comprehensive data collection system that handles:
    - Web scraping with intelligent content extraction
    - PDF extraction with layout preservation
    - Document processing (DOCX, TXT, MD)
    - Metadata extraction and source attribution
    - Quality filtering and validation
    """
    
    def __init__(self, config: Dict[str, Any]):
        """Initialize the data collector with configuration."""
        self.config = config
        self.logger = logging.getLogger(__name__)
        
        # Configuration parameters
        self.output_dir = Path(config.get('output_dir', 'data/raw'))
        self.max_pages_per_source = config.get('max_pages_per_source', 100)
        self.scraping_delay = config.get('scraping_delay', 1.0)
        self.web_sources = config.get('web_sources', [])
        self.pdf_sources = config.get('pdf_sources', [])
        self.timeout = config.get('timeout', 30)
        self.max_retries = config.get('max_retries', 3)
        self.user_agent = config.get('user_agent', 
            'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36')
        
        # Content filtering
        self.min_content_length = config.get('min_content_length', 100)
        self.max_content_length = config.get('max_content_length', 1000000)  # 1MB
        self.allowed_domains = config.get('allowed_domains', [])
        self.blocked_domains = config.get('blocked_domains', [])
        self.blocked_extensions = config.get('blocked_extensions', ['.jpg', '.png', '.gif', '.mp4', '.avi'])
        
        # Create output directory
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # Initialize session for web requests with better configuration
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': self.user_agent,
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
            'Accept-Language': 'en-US,en;q=0.5',
            'Accept-Encoding': 'gzip, deflate',
            'Connection': 'keep-alive',
            'Upgrade-Insecure-Requests': '1',
        })
        
        # Set up request adapter with retry strategy
        from requests.adapters import HTTPAdapter
        from urllib3.util.retry import Retry
        
        retry_strategy = Retry(
            total=self.max_retries,
            backoff_factor=1,
            status_forcelist=[429, 500, 502, 503, 504],
        )
        adapter = HTTPAdapter(max_retries=retry_strategy)
        self.session.mount("http://", adapter)
        self.session.mount("https://", adapter)
        
        # Statistics tracking
        self.stats = {
            'pages_scraped': 0,
            'pages_failed': 0,
            'pdfs_processed': 0,
            'pdfs_failed': 0,
            'total_content_chars': 0,
            'start_time': datetime.now()
        }
        
        self.logger.info(f"DataCollector initialized with output dir: {self.output_dir}")
    
    def scrape_web_sources(self, sources: Optional[List[str]] = None) -> List[Dict[str, Any]]:
        """
        Scrape web sources for domain-specific content with comprehensive metadata extraction.
        """
        if not WEB_SCRAPING_AVAILABLE:
            self.logger.error("Web scraping dependencies not available")
            return []
        
        sources = sources or self.web_sources
        if not sources:
            self.logger.warning("No web sources provided for scraping")
            return []
        
        self.logger.info(f"Starting web scraping for {len(sources)} sources")
        
        scraped_data = []
        
        for i, source_url in enumerate(sources):
            self.logger.info(f"Scraping source {i+1}/{len(sources)}: {source_url}")
            
            try:
                # Validate URL
                if not self._is_valid_url(source_url):
                    self.logger.warning(f"Invalid or blocked URL: {source_url}")
                    continue
                
                # Scrape single source
                source_data = self._scrape_single_source(source_url)
                if source_data:
                    scraped_data.extend(source_data)
                    self.stats['pages_scraped'] += len(source_data)
                
                # Respect rate limiting
                if i < len(sources) - 1:
                    time.sleep(self.scraping_delay)
                    
            except Exception as e:
                self.logger.error(f"Failed to scrape {source_url}: {str(e)}")
                self.stats['pages_failed'] += 1
                continue
        
        self.logger.info(f"Web scraping completed. Collected {len(scraped_data)} pages")
        return scraped_data
    
    def _is_valid_url(self, url: str) -> bool:
        """Validate URL against domain restrictions and blocked extensions."""
        try:
            parsed = urlparse(url)
            domain = parsed.netloc.lower()
            path = parsed.path.lower()
            
            # Check blocked domains
            if self.blocked_domains and any(blocked in domain for blocked in self.blocked_domains):
                return False
            
            # Check allowed domains (if specified)
            if self.allowed_domains and not any(allowed in domain for allowed in self.allowed_domains):
                return False
            
            # Check blocked extensions
            if any(path.endswith(ext) for ext in self.blocked_extensions):
                return False
            
            return True
            
        except Exception:
            return False
    
    def _scrape_single_source(self, url: str) -> List[Dict[str, Any]]:
        """Scrape a single web source and extract content with metadata."""
        pages_data = []
        visited_urls = set()
        urls_to_visit = [url]
        pages_scraped = 0
        base_domain = urlparse(url).netloc
        
        while urls_to_visit and pages_scraped < self.max_pages_per_source:
            current_url = urls_to_visit.pop(0)
            
            if current_url in visited_urls:
                continue
            
            visited_urls.add(current_url)
            
            try:
                # Fetch page content with retries
                response = self._fetch_page(current_url)
                if not response:
                    continue
                
                # Check content type
                content_type = response.headers.get('content-type', '').lower()
                if 'text/html' not in content_type:
                    self.logger.debug(f"Skipping non-HTML content: {current_url}")
                    continue
                
                # Parse HTML content
                soup = BeautifulSoup(response.content, 'html.parser')
                
                # Extract page data
                page_data = self._extract_page_content(soup, current_url, response)
                if page_data:
                    pages_data.append(page_data)
                    pages_scraped += 1
                    self.stats['total_content_chars'] += len(page_data['content'])
                
                # Find additional URLs to scrape (same domain only)
                if pages_scraped < self.max_pages_per_source:
                    new_urls = self._extract_internal_links(soup, current_url, base_domain)
                    for new_url in new_urls[:5]:  # Limit new URLs per page
                        if new_url not in visited_urls and new_url not in urls_to_visit:
                            urls_to_visit.append(new_url)
                
                # Small delay between pages
                time.sleep(0.5)
                
            except Exception as e:
                self.logger.warning(f"Failed to scrape page {current_url}: {str(e)}")
                continue
        
        return pages_data
    
    def _fetch_page(self, url: str) -> Optional[requests.Response]:
        """Fetch a single page with error handling and retries."""
        try:
            response = self.session.get(url, timeout=self.timeout)
            response.raise_for_status()
            
            # Check response size
            content_length = len(response.content)
            if content_length > self.max_content_length:
                self.logger.warning(f"Content too large ({content_length} bytes): {url}")
                return None
            
            return response
            
        except requests.exceptions.RequestException as e:
            self.logger.warning(f"Request failed for {url}: {str(e)}")
            return None
        except Exception as e:
            self.logger.warning(f"Unexpected error fetching {url}: {str(e)}")
            return None
    
    def _extract_page_content(self, soup: BeautifulSoup, url: str, response: requests.Response) -> Optional[Dict[str, Any]]:
        """Extract content and comprehensive metadata from a single page."""
        try:
            # Extract title
            title_tag = soup.find('title')
            title = title_tag.get_text().strip() if title_tag else ''
            
            # Extract main content
            content = self._extract_main_content(soup)
            
            # Apply content filters
            if not self._validate_content(content, url):
                return None
            
            # Extract metadata
            metadata = self._extract_page_metadata(soup, url, response)
            
            # Extract additional structured data
            structured_data = self._extract_structured_data(soup)
            if structured_data:
                metadata['structured_data'] = structured_data
            
            # Create content hash for deduplication
            content_hash = hashlib.md5(content.encode('utf-8')).hexdigest()
            
            # Extract language
            language = self._detect_language(soup, content)
            
            page_data = {
                'url': url,
                'title': title,
                'content': content,
                'content_hash': content_hash,
                'language': language,
                'metadata': metadata,
                'source_type': 'web',
                'collected_at': datetime.now().isoformat(),
                'content_stats': {
                    'char_count': len(content),
                    'word_count': len(content.split()),
                    'paragraph_count': content.count('\n\n') + 1
                }
            }
            
            return page_data
            
        except Exception as e:
            self.logger.warning(f"Failed to extract content from {url}: {str(e)}")
            return None
    
    def _validate_content(self, content: str, url: str) -> bool:
        """Validate content quality and length."""
        if not content or len(content.strip()) < self.min_content_length:
            self.logger.debug(f"Content too short for {url}: {len(content)} chars")
            return False
        
        if len(content) > self.max_content_length:
            self.logger.debug(f"Content too long for {url}: {len(content)} chars")
            return False
        
        # Check for minimal text content (not just whitespace/symbols)
        word_count = len([w for w in content.split() if w.isalpha()])
        if word_count < 10:
            self.logger.debug(f"Insufficient text content for {url}: {word_count} words")
            return False
        
        return True
    
    def _extract_main_content(self, soup: BeautifulSoup) -> str:
        """Extract main text content from HTML, removing navigation and ads."""
        # Remove unwanted elements
        unwanted_tags = ['script', 'style', 'nav', 'header', 'footer', 'aside', 
                        'form', 'button', 'input', 'select', 'textarea']
        unwanted_classes = ['nav', 'navigation', 'menu', 'sidebar', 'footer', 
                           'header', 'ad', 'advertisement', 'popup', 'modal']
        
        for element in soup(unwanted_tags):
            element.decompose()
        
        # Remove elements with unwanted classes
        for class_name in unwanted_classes:
            for element in soup.find_all(class_=re.compile(class_name, re.I)):
                element.decompose()
        
        # Try to find main content areas
        main_content = ''
        
        # Priority 1: Look for semantic HTML5 tags
        main_containers = soup.find_all(['main', 'article'])
        if main_containers:
            for container in main_containers:
                main_content += container.get_text(separator=' ', strip=True) + '\n\n'
        
        # Priority 2: Look for content-specific divs
        if not main_content:
            content_divs = soup.find_all('div', class_=re.compile(
                r'(content|main|article|post|body|text)', re.I))
            if content_divs:
                for div in content_divs:
                    main_content += div.get_text(separator=' ', strip=True) + '\n\n'
        
        # Priority 3: Extract all paragraph and heading text
        if not main_content:
            text_tags = soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'li', 'blockquote'])
            for tag in text_tags:
                text = tag.get_text(strip=True)
                if len(text) > 10:  # Filter out very short text
                    main_content += text + '\n'
        
        # Clean up the content
        main_content = self._clean_text(main_content)
        
        return main_content.strip()
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text content."""
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n', '\n\n', text)  # Remove excessive newlines
        text = re.sub(r'[ \t]+', ' ', text)  # Normalize whitespace
        
        # Remove common web artifacts
        text = re.sub(r'\s+', ' ', text)  # Multiple spaces to single space
        text = re.sub(r'^\s+|\s+$', '', text, flags=re.MULTILINE)  # Trim lines
        
        # Remove repeated punctuation
        text = re.sub(r'[.]{3,}', '...', text)
        text = re.sub(r'[-]{3,}', '---', text)
        
        return text
    
    def _extract_page_metadata(self, soup: BeautifulSoup, url: str, response: requests.Response) -> Dict[str, Any]:
        """Extract comprehensive metadata from page headers and meta tags."""
        metadata = {
            'url': url,
            'domain': urlparse(url).netloc,
            'response_status': response.status_code,
            'content_type': response.headers.get('content-type', ''),
            'content_length': len(response.content),
            'last_modified': response.headers.get('last-modified', ''),
            'server': response.headers.get('server', ''),
            'cache_control': response.headers.get('cache-control', ''),
        }
        
        # Extract meta tags
        meta_tags = {}
        for meta in soup.find_all('meta'):
            # Handle different meta tag formats
            name = (meta.get('name') or meta.get('property') or 
                   meta.get('http-equiv') or meta.get('itemprop'))
            content = meta.get('content')
            if name and content:
                meta_tags[name.lower()] = content
        
        metadata['meta_tags'] = meta_tags
        
        # Extract Open Graph data
        og_data = {}
        for meta in soup.find_all('meta'):
            property_name = meta.get('property', '')
            if property_name.startswith('og:'):
                og_data[property_name] = meta.get('content', '')
        if og_data:
            metadata['open_graph'] = og_data
        
        # Extract Twitter Card data
        twitter_data = {}
        for meta in soup.find_all('meta'):
            name = meta.get('name', '')
            if name.startswith('twitter:'):
                twitter_data[name] = meta.get('content', '')
        if twitter_data:
            metadata['twitter_card'] = twitter_data
        
        # Extract canonical URL
        canonical = soup.find('link', rel='canonical')
        if canonical:
            metadata['canonical_url'] = canonical.get('href')
        
        # Extract page description
        description_meta = soup.find('meta', attrs={'name': 'description'})
        if description_meta:
            metadata['description'] = description_meta.get('content', '')
        
        # Extract keywords
        keywords_meta = soup.find('meta', attrs={'name': 'keywords'})
        if keywords_meta:
            metadata['keywords'] = keywords_meta.get('content', '').split(',')
        
        # Extract author
        author_meta = soup.find('meta', attrs={'name': 'author'})
        if author_meta:
            metadata['author'] = author_meta.get('content', '')
        
        return metadata
    
    def _extract_structured_data(self, soup: BeautifulSoup) -> List[Dict[str, Any]]:
        """Extract structured data (JSON-LD, Microdata) from the page."""
        structured_data = []
        
        # Extract JSON-LD
        json_ld_scripts = soup.find_all('script', type='application/ld+json')
        for script in json_ld_scripts:
            try:
                if script.string:
                    data = json.loads(script.string)
                    structured_data.append({
                        'type': 'json-ld',
                        'data': data
                    })
            except json.JSONDecodeError:
                continue
        
        # Extract Microdata (basic extraction)
        microdata_items = soup.find_all(attrs={'itemtype': True})
        for item in microdata_items:
            try:
                item_data = {
                    'type': 'microdata',
                    'itemtype': item.get('itemtype'),
                    'properties': {}
                }
                
                for prop in item.find_all(attrs={'itemprop': True}):
                    prop_name = prop.get('itemprop')
                    prop_value = prop.get('content') or prop.get_text(strip=True)
                    item_data['properties'][prop_name] = prop_value
                
                if item_data['properties']:
                    structured_data.append(item_data)
            except Exception:
                continue
        
        return structured_data
    
    def _detect_language(self, soup: BeautifulSoup, content: str) -> str:
        """Detect the language of the content."""
        # Check HTML lang attribute
        html_tag = soup.find('html')
        if html_tag and html_tag.get('lang'):
            return html_tag.get('lang')
        
        # Check meta tags
        lang_meta = soup.find('meta', attrs={'http-equiv': 'content-language'})
        if lang_meta:
            return lang_meta.get('content', 'unknown')
        
        # Simple language detection based on common words
        english_indicators = ['the', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by']
        words = content.lower().split()[:100]  # Check first 100 words
        english_count = sum(1 for word in words if word in english_indicators)
        
        if english_count > 5:
            return 'en'
        
        return 'unknown'
    
    def _extract_internal_links(self, soup: BeautifulSoup, base_url: str, base_domain: str) -> List[str]:
        """Extract internal links from the page."""
        internal_links = []
        
        for link in soup.find_all('a', href=True):
            href = link['href']
            
            # Skip javascript, mailto, tel links
            if any(href.startswith(prefix) for prefix in ['javascript:', 'mailto:', 'tel:', '#']):
                continue
            
            # Convert relative URLs to absolute
            absolute_url = urljoin(base_url, href)
            parsed_url = urlparse(absolute_url)
            
            # Check if it's an internal link (same domain)
            if parsed_url.netloc == base_domain:
                # Additional filtering
                if self._is_valid_url(absolute_url):
                    internal_links.append(absolute_url)
        
        return list(set(internal_links))  # Remove duplicates
    
    def extract_pdf_content(self, sources: Optional[List[str]] = None) -> List[Dict[str, Any]]:
        """
        Extract content from PDF files with layout preservation and comprehensive metadata.
        """
        if not PDF_PROCESSING_AVAILABLE:
            self.logger.error("PDF processing dependencies not available")
            return []
        
        sources = sources or self.pdf_sources
        if not sources:
            self.logger.warning("No PDF sources provided")
            return []
        
        self.logger.info(f"Processing {len(sources)} PDF sources")
        
        pdf_data = []
        
        for pdf_source in sources:
            self.logger.info(f"Processing PDF: {pdf_source}")
            
            try:
                # Handle different source types
                if pdf_source.startswith(('http://', 'https://')):
                    # Download PDF from URL
                    pdf_content = self._download_pdf(pdf_source)
                    source_path = pdf_source
                else:
                    # Local PDF file
                    pdf_path = Path(pdf_source)
                    if not pdf_path.exists():
                        self.logger.error(f"PDF file not found: {pdf_source}")
                        continue
                    
                    with open(pdf_path, 'rb') as f:
                        pdf_content = f.read()
                    source_path = str(pdf_path)
                
                # Extract content and metadata
                extracted_data = self._extract_pdf_content(pdf_content, source_path)
                if extracted_data:
                    pdf_data.append(extracted_data)
                    self.stats['pdfs_processed'] += 1
                    self.stats['total_content_chars'] += len(extracted_data['content'])
                    
            except Exception as e:
                self.logger.error(f"Failed to process PDF {pdf_source}: {str(e)}")
                self.stats['pdfs_failed'] += 1
                continue
        
        self.logger.info(f"PDF processing completed. Extracted {len(pdf_data)} documents")
        return pdf_data
    
    def _download_pdf(self, url: str) -> bytes:
        """Download PDF from URL with validation."""
        try:
            response = self.session.get(url, timeout=60)
            response.raise_for_status()
            
            # Verify it's a PDF
            content_type = response.headers.get('content-type', '').lower()
            if 'pdf' not in content_type and not url.lower().endswith('.pdf'):
                # Check if content starts with PDF signature
                if not response.content.startswith(b'%PDF-'):
                    raise ValueError(f"URL does not appear to contain a PDF: {url}")
            
            # Check file size
            if len(response.content) > self.max_content_length:
                raise ValueError(f"PDF file too large: {len(response.content)} bytes")
            
            return response.content
            
        except requests.exceptions.RequestException as e:
            raise Exception(f"Failed to download PDF from {url}: {str(e)}")
    
    def _extract_pdf_content(self, pdf_content: bytes, source_path: str) -> Optional[Dict[str, Any]]:
        """Extract text content and metadata from PDF bytes using multiple methods."""
        try:
            # Try pdfplumber first (better layout preservation)
            extracted_data = None
            
            try:
                with pdfplumber.open(io.BytesIO(pdf_content)) as pdf:
                    extracted_data = self._extract_with_pdfplumber(pdf, source_path)
            except Exception as e:
                self.logger.warning(f"pdfplumber failed for {source_path}: {str(e)}")
                
                # Fallback to PyPDF2
                try:
                    extracted_data = self._extract_with_pypdf2(pdf_content, source_path)
                except Exception as e2:
                    self.logger.error(f"PyPDF2 also failed for {source_path}: {str(e2)}")
                    return None
            
            return extracted_data
            
        except Exception as e:
            self.logger.error(f"Failed to extract PDF content from {source_path}: {str(e)}")
            return None
    
    def _extract_with_pdfplumber(self, pdf, source_path: str) -> Dict[str, Any]:
        """Extract content using pdfplumber (preserves layout better)."""
        pages_content = []
        total_text = ""
        
        for page_num, page in enumerate(pdf.pages):
            try:
                # Extract text with layout
                page_text = page.extract_text()
                if page_text:
                    cleaned_text = self._clean_text(page_text)
                    pages_content.append({
                        'page_number': page_num + 1,
                        'content': cleaned_text.strip(),
                        'char_count': len(cleaned_text),
                        'word_count': len(cleaned_text.split())
                    })
                    total_text += cleaned_text + "\n\n"
                
            except Exception as e:
                self.logger.warning(f"Failed to extract page {page_num + 1} from {source_path}: {str(e)}")
                continue
        
        # Extract metadata
        metadata = {
            'total_pages': len(pdf.pages),
            'pages_extracted': len(pages_content),
            'extraction_method': 'pdfplumber',
            'source_path': source_path,
            'total_chars': len(total_text),
            'total_words': len(total_text.split()),
            'extracted_at': datetime.now().isoformat()
        }
        
        # Try to get PDF metadata
        try:
            if hasattr(pdf, 'metadata') and pdf.metadata:
                pdf_metadata = {}
                for key, value in pdf.metadata.items():
                    try:
                        pdf_metadata[str(key)] = str(value)
                    except:
                        continue
                metadata['pdf_metadata'] = pdf_metadata
        except Exception:
            pass
        
        content_hash = hashlib.md5(total_text.encode('utf-8')).hexdigest()
        
        return {
            'source_path': source_path,
            'content': total_text.strip(),
            'content_hash': content_hash,
            'pages': pages_content,
            'metadata': metadata,
            'source_type': 'pdf',
            'collected_at': datetime.now().isoformat(),
            'content_stats': {
                'char_count': len(total_text),
                'word_count': len(total_text.split()),
                'page_count': len(pages_content)
            }
        }
    
    def extract_document_content(self, file_path: str) -> Optional[Dict[str, Any]]:
        """
        Extract content from various document formats (DOCX, TXT, MD, etc.).
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            self.logger.error(f"File not found: {file_path}")
            return None
        
        self.logger.info(f"Extracting content from: {file_path}")
        
        try:
            file_extension = file_path.suffix.lower()
            
            if file_extension == '.docx':
                return self._extract_docx_content(file_path)
            elif file_extension in ['.txt', '.md', '.markdown']:
                return self._extract_text_content(file_path)
            elif file_extension == '.html':
                return self._extract_html_content(file_path)
            elif file_extension == '.csv':
                return self._extract_csv_content(file_path)
            elif file_extension == '.json':
                return self._extract_json_content(file_path)
            else:
                self.logger.warning(f"Unsupported file format: {file_extension}")
                return None
                
        except Exception as e:
            self.logger.error(f"Failed to extract content from {file_path}: {str(e)}")
            return None
    
    def _extract_docx_content(self, file_path: Path) -> Dict[str, Any]:
        """Extract content from DOCX files."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not available for DOCX processing")
        
        doc = DocxDocument(str(file_path))
        
        # Extract paragraphs
        paragraphs = []
        full_text = ""
        
        for para in doc.paragraphs:
            if para.text.strip():
                cleaned_text = self._clean_text(para.text)
                paragraphs.append(cleaned_text.strip())
                full_text += cleaned_text + "\n"
        
        # Extract tables
        tables_content = []
        for table in doc.tables:
            table_text = ""
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                table_text += row_text + "\n"
            if table_text.strip():
                tables_content.append(table_text.strip())
                full_text += table_text + "\n"
        
        # Extract metadata
        core_props = doc.core_properties
        metadata = {
            'title': core_props.title or '',
            'author': core_props.author or '',
            'subject': core_props.subject or '',
            'keywords': core_props.keywords or '',
            'created': core_props.created.isoformat() if core_props.created else '',
            'modified': core_props.modified.isoformat() if core_props.modified else '',
            'last_modified_by': core_props.last_modified_by or '',
            'paragraphs_count': len(paragraphs),
            'tables_count': len(tables_content),
            'extraction_method': 'python-docx',
            'source_path': str(file_path),
            'total_chars': len(full_text),
            'total_words': len(full_text.split()),
            'extracted_at': datetime.now().isoformat()
        }
        
        # File statistics
        file_stats = file_path.stat()
        metadata.update({
            'file_size': file_stats.st_size,
            'file_modified': datetime.fromtimestamp(file_stats.st_mtime).isoformat(),
            'file_created': datetime.fromtimestamp(file_stats.st_ctime).isoformat()
        })
        
        content_hash = hashlib.md5(full_text.encode('utf-8')).hexdigest()
        
        return {
            'source_path': str(file_path),
            'content': full_text.strip(),
            'content_hash': content_hash,
            'paragraphs': paragraphs,
            'tables': tables_content,
            'metadata': metadata,
            'source_type': 'docx',
            'collected_at': datetime.now().isoformat(),
            'content_stats': {
                'char_count': len(full_text),
                'word_count': len(full_text.split()),
                'paragraph_count': len(paragraphs),
                'table_count': len(tables_content)
            }
        }
    
    def _extract_text_content(self, file_path: Path) -> Dict[str, Any]:
        """Extract content from plain text files with encoding detection."""
        # Try different encodings
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252', 'ascii']
        content = None
        encoding_used = None
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                encoding_used = encoding
                break
            except UnicodeDecodeError:
                continue
            except Exception as e:
                self.logger.warning(f"Error reading {file_path} with {encoding}: {str(e)}")
                continue
        
        if content is None:
            raise ValueError(f"Could not decode file with any supported encoding: {file_path}")
        
        # Clean content
        content = self._clean_text(content)
        
        # Basic file stats
        file_stats = file_path.stat()
        
        # Detect if it's markdown
        is_markdown = file_path.suffix.lower() in ['.md', '.markdown']
        
        metadata = {
            'encoding': encoding_used,
            'file_size': file_stats.st_size,
            'lines_count': len(content.splitlines()),
            'is_markdown': is_markdown,
            'extraction_method': 'text_reader',
            'source_path': str(file_path),
            'total_chars': len(content),
            'total_words': len(content.split()),
            'extracted_at': datetime.now().isoformat(),
            'file_modified': datetime.fromtimestamp(file_stats.st_mtime).isoformat(),
            'file_created': datetime.fromtimestamp(file_stats.st_ctime).isoformat()
        }
        
        content_hash = hashlib.md5(content.encode('utf-8')).hexdigest()
        
        return {
            'source_path': str(file_path),
            'content': content.strip(),
            'content_hash': content_hash,
            'metadata': metadata,
            'source_type': 'text',
            'collected_at': datetime.now().isoformat(),
            'content_stats': {
                'char_count': len(content),
                'word_count': len(content.split()),
                'line_count': len(content.splitlines())
            }
        }
    
    def _extract_html_content(self, file_path: Path) -> Dict[str, Any]:
        """Extract content from local HTML files."""
        if not WEB_SCRAPING_AVAILABLE:
            raise ImportError("BeautifulSoup not available for HTML processing")
        
        with open(file_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Extract title
        title_tag = soup.find('title')
        title = title_tag.get_text().strip() if title_tag else ''
        
        # Extract main content
        content = self._extract_main_content(soup)
        
        # Extract metadata
        metadata = self._extract_page_metadata(soup, str(file_path), None)
        metadata.update({
            'extraction_method': 'html_file',
            'source_path': str(file_path),
            'total_chars': len(content),
            'total_words': len(content.split()),
            'extracted_at': datetime.now().isoformat()
        })
        
        # File statistics
        file_stats = file_path.stat()
        metadata.update({
            'file_size': file_stats.st_size,
            'file_modified': datetime.fromtimestamp(file_stats.st_mtime).isoformat()
        })
        
        content_hash = hashlib.md5(content.encode('utf-8')).hexdigest()
        
        return {
            'source_path': str(file_path),
            'title': title,
            'content': content,
            'content_hash': content_hash,
            'metadata': metadata,
            'source_type': 'html',
            'collected_at': datetime.now().isoformat(),
            'content_stats': {
                'char_count': len(content),
                'word_count': len(content.split())
            }
        }
    
    def _extract_csv_content(self, file_path: Path) -> Dict[str, Any]:
        """Extract content from CSV files."""
        try:
            import pandas as pd
            
            # Read CSV with error handling
            try:
                df = pd.read_csv(file_path)
            except Exception:
                # Try with different encoding
                df = pd.read_csv(file_path, encoding='latin-1')
            
            # Convert to text representation
            content = df.to_string(index=False)
            
            # Extract metadata
            metadata = {
                'rows': len(df),
                'columns': len(df.columns),
                'column_names': df.columns.tolist(),
                'extraction_method': 'pandas_csv',
                'source_path': str(file_path),
                'total_chars': len(content),
                'extracted_at': datetime.now().isoformat()
            }
            
            # File statistics
            file_stats = file_path.stat()
            metadata.update({
                'file_size': file_stats.st_size,
                'file_modified': datetime.fromtimestamp(file_stats.st_mtime).isoformat()
            })
            
            content_hash = hashlib.md5(content.encode('utf-8')).hexdigest()
            
            return {
                'source_path': str(file_path),
                'content': content,
                'content_hash': content_hash,
                'dataframe_info': {
                    'shape': df.shape,
                    'columns': df.columns.tolist(),
                    'dtypes': df.dtypes.to_dict()
                },
                'metadata': metadata,
                'source_type': 'csv',
                'collected_at': datetime.now().isoformat(),
                'content_stats': {
                    'char_count': len(content),
                    'row_count': len(df),
                    'column_count': len(df.columns)
                }
            }
            
        except ImportError:
            # Fallback to basic text processing
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            lines = content.splitlines()
            
            metadata = {
                'lines_count': len(lines),
                'extraction_method': 'text_csv',
                'source_path': str(file_path),
                'total_chars': len(content),
                'extracted_at': datetime.now().isoformat()
            }
            
            content_hash = hashlib.md5(content.encode('utf-8')).hexdigest()
            
            return {
                'source_path': str(file_path),
                'content': content,
                'content_hash': content_hash,
                'metadata': metadata,
                'source_type': 'csv',
                'collected_at': datetime.now().isoformat()
            }
    
    def _extract_json_content(self, file_path: Path) -> Dict[str, Any]:
        """Extract content from JSON files."""
        with open(file_path, 'r', encoding='utf-8') as f:
            json_data = json.load(f)
        
        # Convert JSON to readable text
        content = json.dumps(json_data, indent=2, ensure_ascii=False)
        
        # Analyze JSON structure
        def analyze_json_structure(obj, path=""):
            structure = {}
            if isinstance(obj, dict):
                structure['type'] = 'object'
                structure['keys'] = list(obj.keys())
                structure['key_count'] = len(obj.keys())
            elif isinstance(obj, list):
                structure['type'] = 'array'
                structure['length'] = len(obj)
                if obj:
                    structure['item_type'] = type(obj[0]).__name__
            else:
                structure['type'] = type(obj).__name__
            return structure
        
        json_structure = analyze_json_structure(json_data)
        
        metadata = {
            'json_structure': json_structure,
            'extraction_method': 'json_reader',
            'source_path': str(file_path),
            'total_chars': len(content),
            'extracted_at': datetime.now().isoformat()
        }
        
        # File statistics
        file_stats = file_path.stat()
        metadata.update({
            'file_size': file_stats.st_size,
            'file_modified': datetime.fromtimestamp(file_stats.st_mtime).isoformat()
        })
        
        content_hash = hashlib.md5(content.encode('utf-8')).hexdigest()
        
        return {
            'source_path': str(file_path),
            'content': content,
            'content_hash': content_hash,
            'json_data': json_data,
            'metadata': metadata,
            'source_type': 'json',
            'collected_at': datetime.now().isoformat(),
            'content_stats': {
                'char_count': len(content),
                'json_size': len(str(json_data))
            }
        }
    
    def save_raw_data(self, data: Dict[str, Any], filename: Optional[str] = None) -> str:
        """Save collected raw data to disk with comprehensive metadata."""
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"raw_data_{timestamp}.json"
        
        save_path = self.output_dir / filename
        
        # Calculate statistics
        web_content = data.get('web_content', [])
        pdf_content = data.get('pdf_content', [])
        total_sources = len(web_content) + len(pdf_content)
        
        total_chars = sum(
            len(item.get('content', '')) 
            for items in [web_content, pdf_content]
            for item in items
        )
        
        # Add comprehensive collection summary
        collection_summary = {
            'collection_timestamp': datetime.now().isoformat(),
            'collection_duration': str(datetime.now() - self.stats['start_time']),
            'total_sources': total_sources,
            'web_sources_count': len(web_content),
            'pdf_sources_count': len(pdf_content),
            'total_content_chars': total_chars,
            'average_content_length': total_chars / total_sources if total_sources > 0 else 0,
            'collection_stats': self.stats.copy(),
            'unique_domains': len(set(
                urlparse(item.get('url', '')).netloc 
                for item in web_content 
                if item.get('url')
            )),
            'languages_detected': list(set(
                item.get('language', 'unknown') 
                for item in web_content 
                if item.get('language')
            )),
            'content_types': {
                'web_pages': len(web_content),
                'pdf_documents': len(pdf_content),
                'total_pages': sum(
                    len(item.get('pages', [])) 
                    for item in pdf_content
                )
            }
        }
        
        # Combine data with summary
        save_data = {
            **data,
            'collection_summary': collection_summary,
            'pipeline_version': '1.0.0',
            'collector_config': {
                'max_pages_per_source': self.max_pages_per_source,
                'scraping_delay': self.scraping_delay,
                'min_content_length': self.min_content_length,
                'timeout': self.timeout
            }
        }
        
        # Save to JSON with proper formatting
        with open(save_path, 'w', encoding='utf-8') as f:
            json.dump(save_data, f, indent=2, ensure_ascii=False, default=str)
        
        self.logger.info(f"Raw data saved to: {save_path}")
        self.logger.info(f"Collection summary: {total_sources} sources, {total_chars:,} characters")
        
        return str(save_path)
    
    def load_raw_data(self, filename: str) -> Dict[str, Any]:
        """Load previously collected raw data."""
        load_path = self.output_dir / filename
        
        if not load_path.exists():
            raise FileNotFoundError(f"Raw data file not found: {load_path}")
        
        with open(load_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        self.logger.info(f"Raw data loaded from: {load_path}")
        
        # Log summary if available
        if 'collection_summary' in data:
            summary = data['collection_summary']
            self.logger.info(f"Loaded data summary: {summary.get('total_sources', 0)} sources, "
                           f"{summary.get('total_content_chars', 0):,} characters")
        
        return data
    
    def get_collection_stats(self) -> Dict[str, Any]:
        """Get comprehensive statistics about collected data."""
        stats = {
            'collector_stats': self.stats.copy(),
            'output_directory': str(self.output_dir),
            'raw_data_files': [],
            'total_files': 0,
            'total_size_mb': 0.0,
            'collection_timespan': str(datetime.now() - self.stats['start_time'])
        }
        
        # Scan output directory
        if self.output_dir.exists():
            for file_path in self.output_dir.glob('*.json'):
                file_stats = file_path.stat()
                file_info = {
                    'filename': file_path.name,
                    'size_mb': file_stats.st_size / (1024 * 1024),
                    'modified': datetime.fromtimestamp(file_stats.st_mtime).isoformat(),
                    'size_bytes': file_stats.st_size
                }
                stats['raw_data_files'].append(file_info)
                stats['total_size_mb'] += file_info['size_mb']
            
            stats['total_files'] = len(stats['raw_data_files'])
            
            # Sort by modification time (newest first)
            stats['raw_data_files'].sort(key=lambda x: x['modified'], reverse=True)
        
        return stats
    
    def validate_collected_data(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Validate collected data quality and completeness with detailed analysis."""
        validation_results = {
            'is_valid': True,
            'warnings': [],
            'errors': [],
            'statistics': {},
            'quality_metrics': {},
            'recommendations': []
        }
        
        # Analyze web content
        web_content = data.get('web_content', [])
        web_stats = self._analyze_content_quality(web_content, 'web')
        
        # Analyze PDF content
        pdf_content = data.get('pdf_content', [])
        pdf_stats = self._analyze_content_quality(pdf_content, 'pdf')
        
        # Overall validation
        total_valid_items = web_stats['valid_items'] + pdf_stats['valid_items']
        total_items = len(web_content) + len(pdf_content)
        
        if total_valid_items == 0:
            validation_results['is_valid'] = False
            validation_results['errors'].append("No valid content items found")
        elif total_valid_items < total_items * 0.5:
            validation_results['warnings'].append(f"Low success rate: {total_valid_items}/{total_items} items valid")
        
        total_chars = web_stats['total_chars'] + pdf_stats['total_chars']
        if total_chars < 1000:
            validation_results['warnings'].append("Very low total content volume")
            validation_results['recommendations'].append("Consider adding more data sources")
        
        # Quality metrics
        validation_results['quality_metrics'] = {
            'content_diversity': len(set(
                urlparse(item.get('url', item.get('source_path', ''))).netloc 
                for item in web_content + pdf_content
            )),
            'average_content_length': total_chars / total_valid_items if total_valid_items > 0 else 0,
            'language_distribution': self._analyze_language_distribution(web_content),
            'content_type_distribution': {
                'web_pages': len(web_content),
                'pdf_documents': len(pdf_content)
            }
        }
        
        # Detailed statistics
        validation_results['statistics'] = {
            'total_items': total_items,
            'valid_items': total_valid_items,
            'success_rate': total_valid_items / total_items if total_items > 0 else 0,
            'web_content': web_stats,
            'pdf_content': pdf_stats,
            'total_content_chars': total_chars,
            'unique_sources': len(set(
                item.get('url', item.get('source_path', ''))
                for item in web_content + pdf_content
            ))
        }
        
        # Generate recommendations
        if validation_results['quality_metrics']['content_diversity'] < 3:
            validation_results['recommendations'].append("Consider diversifying data sources across more domains")
        
        if validation_results['statistics']['success_rate'] < 0.8:
            validation_results['recommendations'].append("Review and optimize data collection configuration")
        
        return validation_results
    
    def _analyze_content_quality(self, content_items: List[Dict[str, Any]], content_type: str) -> Dict[str, Any]:
        """Analyze quality metrics for a list of content items."""
        valid_items = 0
        total_chars = 0
        short_items = 0
        long_items = 0
        
        for item in content_items:
            content = item.get('content', '')
            content_length = len(content)
            
            if content_length >= self.min_content_length:
                valid_items += 1
                total_chars += content_length
                
                if content_length < 500:
                    short_items += 1
                elif content_length > 10000:
                    long_items += 1
        
        return {
            'total_items': len(content_items),
            'valid_items': valid_items,
            'invalid_items': len(content_items) - valid_items,
            'total_chars': total_chars,
            'average_length': total_chars / valid_items if valid_items > 0 else 0,
            'short_items': short_items,
            'long_items': long_items,
            'success_rate': valid_items / len(content_items) if content_items else 0
        }
    
    def _analyze_language_distribution(self, web_content: List[Dict[str, Any]]) -> Dict[str, int]:
        """Analyze language distribution in web content."""
        language_counts = {}
        for item in web_content:
            language = item.get('language', 'unknown')
            language_counts[language] = language_counts.get(language, 0) + 1
        return language_counts
    
    def cleanup_temp_files(self):
        """Clean up any temporary files created during collection."""
        temp_dirs = [
            self.output_dir / 'temp',
            self.output_dir / '.cache',
            Path.cwd() / '.temp_downloads'
        ]
        
        for temp_dir in temp_dirs:
            if temp_dir.exists():
                import shutil
                try:
                    shutil.rmtree(temp_dir, ignore_errors=True)
                    self.logger.info(f"Cleaned up temporary directory: {temp_dir}")
                except Exception as e:
                    self.logger.warning(f"Failed to clean up {temp_dir}: {str(e)}")
    
    def export_collection_report(self, output_path: Optional[str] = None) -> str:
        """Export a comprehensive collection report."""
        if output_path is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = self.output_dir / f"collection_report_{timestamp}.json"
        
        report = {
            'report_timestamp': datetime.now().isoformat(),
            'collection_stats': self.get_collection_stats(),
            'configuration': {
                'max_pages_per_source': self.max_pages_per_source,
                'scraping_delay': self.scraping_delay,
                'min_content_length': self.min_content_length,
                'max_content_length': self.max_content_length,
                'timeout': self.timeout,
                'web_sources': self.web_sources,
                'pdf_sources': self.pdf_sources
            },
            'runtime_stats': self.stats,
            'dependencies': {
                'web_scraping_available': WEB_SCRAPING_AVAILABLE,
                'pdf_processing_available': PDF_PROCESSING_AVAILABLE,
                'docx_available': DOCX_AVAILABLE,
                'nltk_available': NLTK_AVAILABLE
            }
        }
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(report, f, indent=2, ensure_ascii=False, default=str)
        
        self.logger.info(f"Collection report exported to: {output_path}")
        return str(output_path)
    
    def __del__(self):
        """Cleanup on destruction."""
        try:
            if hasattr(self, 'session'):
                self.session.close()
            self.cleanup_temp_files()
        except Exception:
            pass  # Ignore cleanup errors during destruction

# Utility functions for data collection
def create_data_collector_from_config(config_path: str) -> DataCollector:
    """Create a DataCollector instance from a configuration file."""
    import yaml
    
    with open(config_path, 'r') as f:
        config = yaml.safe_load(f)
    
    data_collection_config = config.get('data_collection', {})
    return DataCollector(data_collection_config)

def batch_collect_sources(sources: List[str], collector_config: Dict[str, Any]) -> Dict[str, Any]:
    """Collect data from multiple sources using batch processing."""
    collector = DataCollector(collector_config)
    
    web_sources = [s for s in sources if s.startswith(('http://', 'https://'))]
    file_sources = [s for s in sources if not s.startswith(('http://', 'https://'))]
    
    collected_data = {
        'web_content': [],
        'pdf_content': [],
        'other_content': [],
        'metadata': {}
    }
    
    # Collect web sources
    if web_sources:
        web_data = collector.scrape_web_sources(web_sources)
        collected_data['web_content'] = web_data
    
    # Collect file sources
    for file_source in file_sources:
        if file_source.lower().endswith('.pdf'):
            pdf_data = collector.extract_pdf_content([file_source])
            collected_data['pdf_content'].extend(pdf_data)
        else:
            doc_data = collector.extract_document_content(file_source)
            if doc_data:
                collected_data['other_content'].append(doc_data)
    
    return collected_data

if __name__ == "__main__":
    # Example usage
    config = {
        'output_dir': 'data/raw',
        'web_sources': ['https://example.com'],
        'pdf_sources': ['example.pdf'],
        'max_pages_per_source': 10,
        'scraping_delay': 1.0
    }
    
    collector = DataCollector(config)
    
    # Test web scraping
    web_data = collector.scrape_web_sources()
    
    # Test PDF extraction
    pdf_data = collector.extract_pdf_content()
    
    # Combine and save
    all_data = {
        'web_content': web_data,
        'pdf_content': pdf_data
    }
    
    saved_path = collector.save_raw_data(all_data)
    print(f"Data collected and saved to: {saved_path}")
    
    # Generate report
    report_path = collector.export_collection_report()
    print(f"Collection report saved to: {report_path}"),
            'pages': pages_content,
            'metadata': metadata,
            'source_type': 'pdf',
            'collected_at': datetime.now().isoformat(),
            'content_stats': {
                'char_count': len(total_text),
                'word_count': len(total_text.split()),
                'page_count': len(pages_content)
            }
        }
    



    def _extract_with_pypdf2(self, pdf_content: bytes, source_path: str) -> Dict[str, Any]:
        """Fallback extraction using PyPDF2."""
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_content))
        
        pages_content = []
        total_text = ""
        
        for page_num in range(len(pdf_reader.pages)):
            try:
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                
                if page_text:
                    cleaned_text = self._clean_text(page_text)
                    pages_content.append({
                        'page_number': page_num + 1,
                        'content': cleaned_text.strip(),
                        'char_count': len(cleaned_text),
                        'word_count': len(cleaned_text.split())
                    })
                    total_text += cleaned_text + "\n\n"
                    
            except Exception as e:
                self.logger.warning(f"Failed to extract page {page_num + 1} from {source_path}: {str(e)}")
                continue
        
        # Extract metadata
        metadata = {
            'total_pages': len(pdf_reader.pages),
            'pages_extracted': len(pages_content),
            'extraction_method': 'PyPDF2',
            'source_path': source_path,
            'total_chars': len(total_text),
            'total_words': len(total_text.split()),
            'extracted_at': datetime.now().isoformat()
        }
        
        # Try to get PDF metadata
        try:
            if pdf_reader.metadata:
                pdf_metadata = {}
                for key, value in pdf_reader.metadata.items():
                    try:
                        pdf_metadata[str(key)] = str(value)
                    except:
                        continue
                metadata['pdf_metadata'] = pdf_metadata
        except Exception:
            pass
        
        content_hash = hashlib.md5(total_text.encode('utf-8')).hexdigest()
        
        return {
            'source_path': source_path,
            'content': total_text.strip(),
            'content_hash': content_hash}